from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(11, 20, 32)
TEAL = RGBColor(0, 109, 107)
TEXT = RGBColor(45, 55, 72)
MUTED = RGBColor(92, 101, 112)


def set_cell_border(paragraph, color="D1D5DB", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def style_run(run, *, font="Arial", size=11, bold=False, color=TEXT, italic=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    style_run(run, size=10.5, color=TEXT)
    return p


def add_section_heading(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    style_run(run, font="Arial", size=9.5, bold=True, color=NAVY)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_border(p, color="C9A46B", size="8", space="3")
    return p


def add_role(document, period, role, company, location, bullets):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(role)
    style_run(run, font="Cambria", size=13, bold=True, color=NAVY)

    p2 = document.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(3)
    run = p2.add_run(f"{company} | {location} | {period}")
    style_run(run, size=10, color=MUTED)

    for bullet in bullets:
        add_bullet(document, bullet)


def build_resume(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)

    logo = Path("Brand/logo-horizontal-single.png")
    if logo.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(8)
        p_logo.add_run().add_picture(str(logo), width=Inches(5.6))

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(10)
    meta = (
        "Beirut, Lebanon  |  +961 3 008457  |  Charbel.21nasr@icloud.com  |  "
        "charbelnasr.online"
    )
    run = p_meta.add_run(meta)
    style_run(run, size=10, color=MUTED)

    add_section_heading(doc, "Professional Summary")
    summary = (
        "Application Developer with production experience designing, delivering, and "
        "maintaining workflow-heavy enterprise platforms for admissions, academic "
        "operations, and internal coordination. Strongest in ASP.NET Core, SQL Server, "
        "Entity Framework, React, and backend-driven systems where operational clarity, "
        "security, and reliability matter."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(summary)
    style_run(run, size=10.5, color=TEXT)

    add_section_heading(doc, "Core Capabilities")
    capabilities = [
        "Enterprise application development across applicant-facing, evaluator-facing, and administrative workflows",
        "ASP.NET Core, SQL Server, Entity Framework, REST APIs, stored procedures, and structured backend services",
        "Workflow design for forms, document handling, recommendation flows, scoring logic, and review-state management",
        "Production support across deployment, issue resolution, permissions, and system reliability",
    ]
    for cap in capabilities:
        add_bullet(doc, cap)

    add_section_heading(doc, "Professional Experience")
    add_role(
        doc,
        "February 2024 - Present",
        "Application Developer",
        "Saint George University of Beirut",
        "Achrafieh, Lebanon",
        [
            "Build and maintain production admissions and academic platforms supporting applicant submission, evaluator review, and administrative operations.",
            "Design and implement workflow-heavy systems covering dynamic forms, document intake, recommendation handling, scoring flows, and internal dashboards.",
            "Develop backend services and APIs using ASP.NET Core, SQL Server, Entity Framework, and secure business logic patterns.",
            "Support day-to-day platform reliability through deployment coordination, issue resolution, permissions handling, and post-launch improvements.",
        ],
    )

    add_role(
        doc,
        "February 2023 - February 2024",
        "IT Support and Junior Development",
        "Saint George University of Beirut",
        "Achrafieh, Lebanon",
        [
            "Supported faculty and staff operations while contributing to internal system maintenance, testing, and technical documentation.",
            "Built early experience with production environments, enterprise processes, and structured technical delivery inside an academic institution.",
        ],
    )

    add_role(
        doc,
        "October 2023 - January 2024",
        "IT Support and Development",
        "Société Sakr",
        "Jounieh, Lebanon",
        [
            "Managed cloud-based storage and internal network resources while supporting business-side technical operations.",
            "Led delivery of an inventory management system from requirements through deployment and user onboarding.",
        ],
    )

    add_section_heading(doc, "Selected Projects")
    projects = [
        "SGUB Undergraduate Admissions Platform: production admissions platform supporting applicant submission, document review, scoring workflows, and administrative decision support.",
        "PGME Residency and Fellowship Applications: multi-step workflow built for document-heavy submissions, structured validation, reviewer coordination, and reliable processing.",
        "Lebanese Red Cross Youth Hub: internal coordination platform with role-based workflows, topic distribution, and operational ownership.",
        "UNISHELF Commerce Experience: commercial product platform combining structured catalog modeling, customer-facing filtering, and cloud-hosted media delivery.",
    ]
    for item in projects:
        add_bullet(doc, item)

    add_section_heading(doc, "Technical Skills")
    skills_lines = [
        ("Languages", "C#, JavaScript, TypeScript, Python, Java, SQL"),
        ("Backend", "ASP.NET Core, .NET MVC, REST APIs, Entity Framework"),
        ("Frontend", "React, Razor Views, HTML5, CSS, JavaScript"),
        ("Data", "SQL Server, Stored Procedures, Triggers, Query Optimization, MongoDB"),
        ("Deployment", "IIS, AWS S3, CloudFront, CI/CD-aware workflows, Docker (basic)"),
        ("Architecture", "Service layers, DTO-based design, role-based access control, server-side validation"),
    ]
    for label, value in skills_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{label}: ")
        style_run(run, size=10.5, bold=True, color=NAVY)
        run = p.add_run(value)
        style_run(run, size=10.5, color=TEXT)

    add_section_heading(doc, "Education and Certification")
    edu_items = [
        "B.S. in Information Technology (AI Emphasis) — Saint George University of Beirut, 2022 - 2025",
        "M.S. in Data Science (In Progress) — Lebanese American University, 2025 - Present",
        "Cisco Certified Support Technician (CCST IT Support)",
    ]
    for item in edu_items:
        add_bullet(doc, item)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    build_resume(Path("Brand/Charbel-Nasr-Resume-Branded.docx"))
